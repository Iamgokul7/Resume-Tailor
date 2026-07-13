"""
test_pipeline.py — Tests the updated ResumeTailor pipeline (file upload, text extraction, fabrication check, and rendering).
Run with: .\\venv\\Scripts\\python.exe test_pipeline.py
"""

import json
import sys
import io
import re
from pathlib import Path

# Insert current directory to path
sys.path.insert(0, str(Path(__file__).parent))

# ── Mock tailored resume (what Gemini would return) ──────────────────────
MOCK_GEMINI_OUTPUT = {
    "summary": (
        "Software Engineer with hands-on experience designing and shipping scalable "
        "backend services, REST APIs, and microservices using Python, Java, and Go. "
        "Comfortable with the full development lifecycle from architecture to CI/CD "
        "deployment. Passionate about writing clean, well-tested code and collaborating "
        "in Agile teams."
    ),
    "skills_selected": [
        {"category": "Languages", "items": ["Python", "JavaScript", "Go", "SQL", "Bash"]},
        {"category": "Frameworks & Libraries", "items": ["FastAPI", "Django", "React", "Pytest"]},
        {"category": "Databases", "items": ["PostgreSQL", "MongoDB", "Redis"]},
        {"category": "Cloud & DevOps", "items": ["AWS (EC2, S3, Lambda)", "Docker", "Kubernetes", "GitHub Actions", "CI/CD"]},
        {"category": "Tools & Practices", "items": ["Git", "REST APIs", "Agile/Scrum", "Postman", "Swagger/OpenAPI"]},
    ],
    "flagship_projects_selected": [
        {
            "name": "SecureVault – Encrypted Secret Manager",
            "dates": "Jan 2024 – Mar 2024",
            "tech_stack": "Python, FastAPI, PostgreSQL, Redis, Docker, JWT, AES-256",
            "github_link": "https://github.com/alexjordan/securevault",
            "bullets": [
                "Designed RESTful FastAPI service with JWT-authenticated endpoints, rate limiting via Redis, and Swagger documentation for all 18 API routes.",
                "Containerized the full stack with Docker Compose; wrote GitHub Actions CI/CD pipeline that runs Pytest suite and auto-deploys to a staging environment on every merge to main.",
                "Built a self-hosted secrets manager with AES-256 encryption, role-based access control, and audit logging, reducing credential exposure risk by enforcing zero-plaintext storage.",
            ],
        }
    ],
    "other_projects_selected": [
        {
            "name": "AutoTest – AI-Assisted QA Framework",
            "dates": "May 2023 – Aug 2023",
            "tech_stack": "Python, Playwright, Selenium, Pytest, Jenkins, GitHub Actions",
            "github_link": "https://github.com/alexjordan/autotest",
            "bullets": [
                "Developed an end-to-end test automation framework using Playwright and Pytest covering 200+ test cases, reducing manual QA cycle time from 3 days to under 4 hours.",
                "Configured Jenkins pipelines and GitHub Actions workflows to run the full Playwright suite on every pull request, with Slack notifications for test failures.",
            ],
        },
    ],
    "internships_selected": [
        {
            "company": "TechCorp Inc.",
            "role": "Software Engineering Intern",
            "dates": "Jun 2022 – Aug 2022",
            "bullets": [
                "Developed and shipped 3 new REST API endpoints in the core Java Spring Boot microservice, adding paginated search and filtering capabilities used by the mobile team.",
                "Optimized a slow PostgreSQL query that ran on the nightly batch job by adding composite indexes and rewriting the JOIN logic, reducing execution time from 8 minutes to under 40 seconds.",
                "Wrote unit and integration tests with JUnit and Mockito, raising code coverage from 61% to 84% for the payments service module.",
            ],
        }
    ],
    "education": [
        {
            "institution": "University of California, Berkeley",
            "degree": "Bachelor of Science in Computer Science",
            "dates": "Aug 2019 – May 2023",
            "gpa": "3.8/4.0",
            "relevant_coursework": "Data Structures, Algorithms, Operating Systems, Computer Networks, Software Engineering",
        }
    ],
    "certifications": [
        {"name": "AWS Certified Developer – Associate", "issuer": "Amazon Web Services", "date": "March 2024"},
        {"name": "Google Associate Cloud Engineer", "issuer": "Google Cloud", "date": "November 2023"},
    ],
}

# ── Dynamic mock master resume text representing the raw text layer ─────
MOCK_MASTER_TEXT = """
Alex Jordan
alex.jordan@email.com | +1 (555) 123-4567 | San Francisco, CA
GitHub: https://github.com/alexjordan | Portfolio: https://alexjordan.dev | LinkedIn: https://linkedin.com/in/alexjordan

SUMMARY OPTIONS:
Software Engineer with hands-on experience designing and shipping scalable backend services, REST APIs, and microservices using Python, Java, and Go. Comfortable with the full development lifecycle from architecture to CI/CD deployment. Passionate about writing clean, well-tested code and collaborating in Agile teams.

TECHNICAL SKILLS:
Languages: Python, JavaScript, TypeScript, Java, Go, SQL, Bash, C++
Frameworks & Libraries: FastAPI, Django, Flask, React, Node.js, Express.js, Spring Boot, Pytest, Selenium, Playwright
Databases: PostgreSQL, MySQL, MongoDB, Redis, SQLite, Elasticsearch
Cloud & DevOps: AWS (EC2, S3, Lambda, RDS, CloudWatch), GCP (Cloud Run, BigQuery, Pub/Sub), Docker, Kubernetes, Terraform, GitHub Actions, Jenkins, CI/CD
Tools & Practices: Git, Linux, REST APIs, GraphQL, Agile/Scrum, JIRA, Postman, Swagger/OpenAPI, Microservices

WORK EXPERIENCE:
TechCorp Inc. - Software Engineering Intern (Jun 2022 – Aug 2022)
* Developed and shipped 3 new REST API endpoints in the core Java Spring Boot microservice, adding paginated search and filtering capabilities used by the mobile team.
* Optimized a slow PostgreSQL query that ran on the nightly batch job by adding composite indexes and rewriting the JOIN logic, reducing execution time from 8 minutes to under 40 seconds.
* Wrote unit and integration tests with JUnit and Mockito, raising code coverage from 61% to 84% for the payments service module.

PROJECTS:
SecureVault – Encrypted Secret Manager (Jan 2024 – Mar 2024)
Tech: Python, FastAPI, PostgreSQL, Redis, Docker, JWT, AES-256
GitHub: https://github.com/alexjordan/securevault
* Built a self-hosted secrets manager with AES-256 encryption, role-based access control, and audit logging, reducing credential exposure risk by enforcing zero-plaintext storage.
* Designed RESTful FastAPI service with JWT-authenticated endpoints, rate limiting via Redis, and Swagger documentation for all 18 API routes.
* Containerized the full stack with Docker Compose; wrote GitHub Actions CI/CD pipeline that runs Pytest suite and auto-deploys to a staging environment on every merge to main.

AutoTest – AI-Assisted QA Framework (May 2023 – Aug 2023)
Tech: Python, Playwright, Selenium, Pytest, Jenkins, GitHub Actions
GitHub: https://github.com/alexjordan/autotest
* Developed an end-to-end test automation framework using Playwright and Pytest covering 200+ test cases, reducing manual QA cycle time from 3 days to under 4 hours.
* Configured Jenkins pipelines and GitHub Actions workflows to run the full Playwright suite on every pull request, with Slack notifications for test failures.

EDUCATION:
University of California, Berkeley - Bachelor of Science in Computer Science (Aug 2019 – May 2023) - GPA: 3.8/4.0
Coursework: Data Structures, Algorithms, Operating Systems, Computer Networks, Software Engineering

CERTIFICATIONS:
AWS Certified Developer – Associate - Amazon Web Services (March 2024)
Google Associate Cloud Engineer - Google Cloud (November 2023)
"""

# Build mock bank for safety checks
bank = {
    "contact": {
        "name": "Alex Jordan",
        "email": "alex.jordan@email.com",
        "phone": "+1 (555) 123-4567",
        "location": "San Francisco, CA"
    },
    "raw_text": MOCK_MASTER_TEXT
}

# Helper to generate in-memory dummy files
def make_dummy_pdf() -> bytes:
    from reportlab.pdfgen import canvas
    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.drawString(100, 750, "Alex Jordan")
    c.drawString(100, 730, "alex.jordan@email.com | +1 (555) 123-4567")
    c.drawString(100, 710, "Python, FastAPI, PostgreSQL, Redis, Docker, CI/CD")
    c.save()
    return buf.getvalue()

def make_dummy_docx() -> bytes:
    import docx
    doc = docx.Document()
    doc.add_paragraph("Alex Jordan")
    doc.add_paragraph("alex.jordan@email.com | +1 (555) 123-4567")
    doc.add_paragraph("Python, FastAPI, PostgreSQL, Redis, Docker, CI/CD")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Test 1: File text extraction ──────────────────────────────────────────
print("\n" + "="*60)
print("TEST 1: File text extraction (_extract_text_from_file)")
print("="*60)

from main import _extract_text_from_file

# Test PDF Extraction
pdf_bytes = make_dummy_pdf()
extracted_pdf_text = _extract_text_from_file("my_resume.pdf", pdf_bytes)
print(f"[OK] Extracted PDF text length: {len(extracted_pdf_text)} characters")
assert "Alex Jordan" in extracted_pdf_text, "Failed to extract name from PDF"
assert "FastAPI" in extracted_pdf_text, "Failed to extract skills from PDF"

# Test DOCX Extraction
docx_bytes = make_dummy_docx()
extracted_docx_text = _extract_text_from_file("my_resume.docx", docx_bytes)
print(f"[OK] Extracted DOCX text length: {len(extracted_docx_text)} characters")
assert "Alex Jordan" in extracted_docx_text, "Failed to extract name from DOCX"
assert "FastAPI" in extracted_docx_text, "Failed to extract skills from DOCX"

# Test Scanned PDF/Empty PDF error handling
try:
    _extract_text_from_file("empty.pdf", b"%PDF-1.4 ... empty bytes ...")
    assert False, "Should have raised a ValueError for bad/scanned PDF"
except ValueError as exc:
    print(f"[OK] Caught expected error for scanned/bad PDF: '{exc}'")
    assert "please make sure it's a text-based PDF" in str(exc)


# ── Test 2: Fabrication safety check ─────────────────────────────────────
print("\n" + "="*60)
print("TEST 2: Fabrication safety check")
print("="*60)

from main import run_fabrication_check

# Should produce ZERO warnings (all tokens come from the mock raw text)
warnings = run_fabrication_check(MOCK_GEMINI_OUTPUT, bank)
print(f"[OK] Clean response: {len(warnings)} fabrication warnings (expected 0)")
assert len(warnings) == 0, f"Unexpected warnings: {warnings}"

# Now inject a fake skill and verify it IS caught
MOCK_WITH_FAKE = json.loads(json.dumps(MOCK_GEMINI_OUTPUT))
MOCK_WITH_FAKE["skills_selected"].append({
    "category": "Fake Skills",
    "items": ["QuantumNeuralSynapse", "HolobaseDB"],
})
warnings_with_fake = run_fabrication_check(MOCK_WITH_FAKE, bank)
print(f"[OK] Injected 2 fake skills -> {len(warnings_with_fake)} warning(s) caught: {warnings_with_fake}")
assert len(warnings_with_fake) == 2, f"Fabrication check FAILED to catch fake skills! Warnings: {warnings_with_fake}"


# ── Test 3: PDF rendering ─────────────────────────────────────────────────
print("\n" + "="*60)
print("TEST 3: PDF rendering (xhtml2pdf)")
print("="*60)

from pdf_service import render_pdf
pdf_path = render_pdf(MOCK_GEMINI_OUTPUT, bank["contact"])
assert pdf_path.exists(), f"PDF not created at {pdf_path}"
size_kb = pdf_path.stat().st_size / 1024
print(f"[OK] PDF created: {pdf_path.name}  ({size_kb:.1f} KB)")
assert size_kb > 5, f"PDF suspiciously small: {size_kb:.1f} KB"


# ── Test 4: Download path validation ─────────────────────────────────────
print("\n" + "="*60)
print("TEST 4: Download path validation")
print("="*60)

filename = pdf_path.name
assert re.match(r"^resume_\d{8}_\d{6}\.pdf$", filename), f"Bad filename: {filename}"
print(f"[OK] Filename format OK: {filename}")


# ── Test 5: Model Fallback Simulation ──────────────────────────────────────
print("\n" + "="*60)
print("TEST 5: Model Fallback Simulation")
print("="*60)

import gemini_service
from gemini_service import analyze_jd_match

# Temporarily prepend an invalid model name to simulate a 404/failure
original_models = list(gemini_service.GEMINI_MODELS)
gemini_service.GEMINI_MODELS = ["gemini-invalid-model-999"] + original_models
print(f"Configured models with fake primary model: {gemini_service.GEMINI_MODELS}")

try:
    # Run a simple match analysis (real API call)
    result = analyze_jd_match(
        master_resume="Alex Jordan\nPython, FastAPI, Docker",
        jd_text="Looking for a Python developer with FastAPI and Docker experience."
    )
    print(f"[OK] Fallback successful! Match Score: {result.get('overall_match_percentage')}%")
    assert "overall_match_percentage" in result, "Fallback result missing match percentage"
except Exception as exc:
    print(f"[FAIL] Fallback test failed: {exc}")
    raise
finally:
    # Restore original models
    gemini_service.GEMINI_MODELS = original_models


# ── Test 6: Quality Gate & Section Regeneration ──────────────────────────
print("\n" + "="*60)
print("TEST 6: Quality Gate & Section Regeneration")
print("="*60)

from gemini_service import check_summary_text, check_bullet_text, validate_and_correct_resume

# Test buzzword checking
bad_summary = "Forward-thinking and passionate developer with experience."
errors = check_summary_text(bad_summary)
print(f"[OK] Buzzword check caught: {errors}")
assert len(errors) > 0, "Failed to catch buzzword opener"

# Test weak opener and vague ending check
bad_bullet = "Worked on database design mirroring project-level reporting requirements."
errors = check_bullet_text(bad_bullet)
print(f"[OK] Bullet checks caught: {errors}")
assert len(errors) > 0, "Failed to catch weak opener or vague ending"

# Run full validate_and_correct_resume online to verify correction works
mock_bad_tailored = {
    "summary": "Forward-thinking and passionate engineer.",
    "flagship_projects_selected": [
        {
            "name": "Project A",
            "bullets": ["Worked on implementing the database to meet requirements."]
        }
    ]
}

final_tailored, warnings = validate_and_correct_resume(
    mock_bad_tailored,
    master_resume="Alex Jordan. Python, database.",
    jd_text="Looking for a database engineer."
)
print(f"[OK] Quality Gate completed. Final Summary: '{final_tailored.get('summary')}'")
print(f"Warnings produced: {warnings}")
assert not check_summary_text(final_tailored.get('summary')), "Regenerated summary still contains buzzwords"


# ── Summary ───────────────────────────────────────────────────────────────
print("\n" + "="*60)
print("ALL TESTS PASSED [OK]")
print("="*60)
print(f"\nGenerated PDF: {pdf_path}")
print("The file upload -> text extraction -> safety check -> PDF pipeline -> quality gate works correctly.")
print("To test the full upload web interface, run uvicorn and open http://localhost:8000 in your browser.\n")
